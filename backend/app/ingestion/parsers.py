"""Document parser dispatch by MIME type.

Each parser returns a list of (page_number, content, content_kind, image_bytes?) tuples
that the pipeline converts to chunks or VLM-extracts.
"""
from __future__ import annotations

import io
import mimetypes
from dataclasses import dataclass
from typing import Literal

ContentKind = Literal["text", "image"]


@dataclass
class ParsedUnit:
    page_number: int
    kind: ContentKind
    text: str = ""
    image_bytes: bytes | None = None
    image_mime: str | None = None


def guess_mime(filename: str, provided: str | None = None) -> str:
    if provided and provided != "application/octet-stream":
        return provided
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "application/octet-stream"


def parse(filename: str, data: bytes, mime: str | None = None) -> list[ParsedUnit]:
    mime = guess_mime(filename, mime)
    if mime in {"text/plain", "text/markdown", "text/csv"}:
        return [ParsedUnit(page_number=1, kind="text", text=data.decode("utf-8", errors="replace"))]
    if mime == "application/pdf":
        return _parse_pdf(data)
    if mime in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        return _parse_docx(data)
    if mime in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }:
        return _parse_xlsx(data)
    if mime.startswith("image/"):
        return [ParsedUnit(page_number=1, kind="image", image_bytes=data, image_mime=mime)]
    # Unknown: treat as UTF-8 text best-effort.
    return [ParsedUnit(page_number=1, kind="text", text=data.decode("utf-8", errors="replace"))]


def _parse_pdf(data: bytes) -> list[ParsedUnit]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    units: list[ParsedUnit] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        units.append(ParsedUnit(page_number=idx, kind="text", text=text))
    # Note: for image-heavy / scanned PDFs, a production path would rasterize each page
    # (pdf2image) and feed it through the VLM. That's the intended integration point
    # — it's plumbed but left out here to keep the parser pure-Python.
    return units


def _parse_docx(data: bytes) -> list[ParsedUnit]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(paragraphs)
    return [ParsedUnit(page_number=1, kind="text", text=text)]


def _parse_xlsx(data: bytes) -> list[ParsedUnit]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    units: list[ParsedUnit] = []
    for sheet_idx, ws in enumerate(wb.worksheets, start=1):
        rows_text: list[str] = [f"# Sheet: {ws.title}"]
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows_text.append(
                    " | ".join("" if c is None else str(c) for c in row)
                )
        units.append(
            ParsedUnit(page_number=sheet_idx, kind="text", text="\n".join(rows_text))
        )
    return units


__all__ = ["ContentKind", "ParsedUnit", "guess_mime", "parse"]
